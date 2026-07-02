"""
Sinh báo cáo Word (.docx) cho dự án 45-class DNA Contributor Identification.
Chạy: python generate_report.py
"""

from pathlib import Path
from datetime import datetime
import json, textwrap

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
OUT  = ROOT / "reports"
OUT.mkdir(exist_ok=True)

# ── Helpers ────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False,
             italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            run.font.size = Pt(14)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(13)
            run.font.bold = True
        else:
            run.font.size = Pt(12)
            run.font.bold = True
    return p


def add_para(doc, text, indent=False, bold_parts=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run(text)
    set_font(run)
    return p


def add_code(doc, code_text):
    """Code block với nền xám, font Courier New."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    # Nền xám
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
            if c_idx > 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


# ── Build document ─────────────────────────────────────────────────────────

def build(doc: Document):

    # ── Title page ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("BÁO CÁO KỸ THUẬT")
    set_font(r, size=16, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Nhận dạng cá nhân trong hỗn hợp DNA:\n"
        "Bộ phân loại đa nhãn 45 lớp trên dữ liệu PROVEDIt"
    )
    set_font(r, size=14, bold=True)

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta_p.add_run(f"Ngày: {datetime.today().strftime('%d/%m/%Y')}")
    set_font(r, size=12, italic=True)

    doc.add_page_break()

    # ── 1. Giới thiệu ───────────────────────────────────────────────────────
    add_heading(doc, "1. Giới thiệu", 1)
    add_para(doc,
        "Phân tích DNA hỗn hợp (DNA mixture) là bài toán quan trọng trong "
        "pháp y — khi mẫu vật chứa vật liệu di truyền của nhiều cá nhân, "
        "hệ thống cần xác định có bao nhiêu người đóng góp và cụ thể là ai. "
        "Các hệ thống truyền thống như STRmix hay EuroForMix dựa trên mô hình "
        "xác suất phức tạp và đòi hỏi chuyên gia vận hành. Gần đây, các "
        "phương pháp học máy bắt đầu cho thấy tiềm năng tự động hóa quá "
        "trình này.", indent=True)

    add_para(doc,
        "Báo cáo này trình bày phần cá nhân của dự án nhóm: xây dựng bộ phân "
        "loại đa nhãn (multi-label classifier) 45 lớp xác định những người "
        "đóng góp đã biết trong một hỗn hợp DNA. Với mỗi mẫu đầu vào là "
        "vector đặc trưng 252 chiều mô tả chiều cao đỉnh (peak height) tại "
        "các vị trí STR loci, đầu ra là vector nhị phân 45 chiều cho biết "
        "donor nào có mặt trong hỗn hợp.", indent=True)

    add_para(doc,
        "Phần OOD detection (phát hiện cá nhân chưa biết) là phần của thành "
        "viên khác trong nhóm (manhngvu) và sẽ được tích hợp vào pipeline "
        "chung ở giai đoạn sau.", indent=True)

    # ── 2. Bài toán ─────────────────────────────────────────────────────────
    add_heading(doc, "2. Định nghĩa bài toán", 1)
    add_para(doc,
        "Cho một mẫu hỗn hợp DNA ký hiệu x ∈ ℝ^252, tìm vector nhãn "
        "y ∈ {0,1}^45 sao cho y_i = 1 khi và chỉ khi donor i là một trong "
        "những người đóng góp vào hỗn hợp. "
        "Đây là bài toán multi-label classification — mỗi mẫu có thể có nhiều "
        "nhãn dương cùng lúc (tương ứng với hỗn hợp 2–5 người).", indent=True)

    add_para(doc,
        "Thiết lập closed-set: toàn bộ quá trình huấn luyện và đánh giá "
        "trong báo cáo này chỉ xét các mẫu mà tất cả donor đều nằm trong "
        "tập 45 người đã biết. Bộ 5 donor hold-out đóng vai 'unknown' và "
        "được xử lý riêng bởi module OOD.", indent=True)

    # ── 3. Dataset ──────────────────────────────────────────────────────────
    add_heading(doc, "3. Dữ liệu", 1)

    add_heading(doc, "3.1 Nguồn dữ liệu", 2)
    add_para(doc,
        "Dữ liệu được lấy từ bộ dữ liệu PROVEDIt (Proficiency Review of "
        "Evidentiary DNA Interpretation Tools), phiên bản RD14-0003. "
        "PROVEDIt là bộ dữ liệu mở do Rutgers University tạo lập, bao gồm "
        "hàng nghìn mẫu STR (Short Tandem Repeat) được phân tích dưới nhiều "
        "điều kiện thí nghiệm khác nhau (loại máy, bộ kit, nồng độ DNA, "
        "điều kiện tiêm).", indent=True)

    add_para(doc,
        "Phiên bản sử dụng trong dự án này: điều kiện GF (GlobalFiler kit), "
        "thời gian tiêm 5 giây, máy 3500 với marker SE33F. Tổng cộng 50 donor "
        "duy nhất (IDs 1–50), bao gồm các mẫu đơn (1-person) và hỗn hợp "
        "(2–5 người).", indent=True)

    add_heading(doc, "3.2 Format file thô", 2)
    add_para(doc,
        "Mỗi file CSV thô chứa kết quả phân tích STR của nhiều mẫu. Mỗi hàng "
        "tương ứng với một cặp (mẫu, locus). Các cột chính:", indent=True)

    add_table(doc,
        ["Cột", "Kiểu", "Mô tả"],
        [
            ["Sample File", "string", "Tên file .hid, chứa mã hóa donor ID và điều kiện thí nghiệm"],
            ["Marker", "string", "Tên STR locus (D3S1358, vWA, D16S539, ...)"],
            ["Dye", "string", "Màu thuốc nhuộm điện di (B, G, Y, R)"],
            ["Allele 1..100", "string", "Giá trị allele (số nguyên/thập phân/OL/NaN)"],
            ["Size 1..100", "float", "Kích thước fragment (bp)"],
            ["Height 1..100", "float", "Chiều cao đỉnh RFU (relative fluorescence unit)"],
        ],
        col_widths=[3.5, 2, 10])

    doc.add_paragraph()
    add_heading(doc, "3.3 Quy ước đặt tên file (naming convention)", 2)
    add_para(doc,
        "Tên file mã hóa toàn bộ thông tin cần thiết để xác định donor. "
        "Ví dụ:", indent=True)

    add_code(doc,
        "1-person : A02_RD14-0003-15d2U60-0.25GF-Q4.5_01.5sec.hid\n"
        "           └─ donor ID = 15 (hai chữ số đầu trước ký tự chữ)\n\n"
        "2-person : A02_RD14-0003-31_32-1;1-M2c-0.03GF-Q2.0_01.5sec.hid\n"
        "           └─ donors = {31, 32} (phân tách bằng dấu _)\n\n"
        "5-person : A02_RD14-0003-43_44_45_46_47-1;1;2;1;1-M4c-...\n"
        "           └─ donors = {43, 44, 45, 46, 47}")

    add_para(doc,
        "Hàm parse được cài đặt trong data/prepare_data.py:", indent=True)

    add_code(doc,
        "def parse_donors(filename: str) -> list[int]:\n"
        "    parts = filename.split('-')\n"
        "    contrib_part = parts[2]          # '31_32' hoặc '15d2U60'\n"
        "    if '_' in contrib_part:\n"
        "        return [int(re.match(r'^(\\d+)', s).group(1))\n"
        "                for s in contrib_part.split('_')]\n"
        "    else:\n"
        "        m = re.match(r'^(\\d+)', contrib_part)\n"
        "        return [int(m.group(1))] if m else []")

    add_heading(doc, "3.4 Tiền xử lý dữ liệu", 2)
    add_para(doc,
        "Pipeline tiền xử lý kế thừa từ module NOC_DNA (bài toán trước của "
        "nhóm), thực hiện các bước sau:", indent=True)

    add_table(doc,
        ["Bước", "Mô tả"],
        [
            ["1. Melt wide → long", "Biến ma trận Allele/Size/Height thành dạng long (mỗi peak một hàng)"],
            ["2. Lọc ngưỡng", "Loại bỏ peaks có Height < 50 RFU (nhiễu), Allele = OL (off-ladder)"],
            ["3. Stutter filter", "Loại stutter artifact (peak N-1, < 15% chiều cao peak chính)"],
            ["4. Pivot genotype", "Pivot về dạng ma trận thưa: Sample File × (Locus, Allele) → Height"],
            ["5. Flatten header", "Đổi tên cột (Marker, Allele) → 'Marker_Allele' (ví dụ: D3S1358_15)"],
        ],
        col_widths=[4, 12])

    doc.add_paragraph()
    add_para(doc,
        "Kết quả: file combined_preprocessed_data.csv với 3.342 mẫu, "
        "252 cột đặc trưng (allele heights theo locus_allele), 1 cột "
        "Sample File và 1 cột target_noc (không dùng trong task này).", indent=True)

    add_heading(doc, "3.5 Phân chia donor và xây dựng nhãn", 2)
    add_para(doc,
        "50 donor được chia thành 45 known và 5 unknown dùng numpy random "
        "seed=42 để đảm bảo tái lập được:", indent=True)

    add_code(doc,
        "rng = np.random.default_rng(42)\n"
        "_shuffled = rng.permutation(list(range(1, 51))).tolist()\n"
        "UNKNOWN_DONORS = sorted(_shuffled[:5])   # [6, 21, 26, 40, 50]\n"
        "KNOWN_DONORS   = sorted(_shuffled[5:])   # 45 donors còn lại")

    add_para(doc,
        "Nhãn mỗi mẫu là vector 45-dim binary: y[i] = 1 khi KNOWN_DONORS[i] "
        "có mặt trong hỗn hợp. Các mẫu chứa bất kỳ donor unknown nào được "
        "tách riêng (open-set, 451 mẫu) để tích hợp OOD sau. "
        "Closed-set còn lại: 2.891 mẫu.", indent=True)

    add_table(doc,
        ["Tập", "Số mẫu", "Tỷ lệ"],
        [
            ["Train", "2.023", "70%"],
            ["Validation", "434", "15%"],
            ["Test", "434", "15%"],
            ["Open-set (OOD)", "451", "—"],
        ],
        col_widths=[4, 3, 3])

    doc.add_page_break()

    # ── 4. Kiến trúc mô hình ────────────────────────────────────────────────
    add_heading(doc, "4. Kiến trúc mô hình", 1)

    add_heading(doc, "4.1 Baseline: XGBoost One-vs-Rest", 2)
    add_para(doc,
        "Baseline sử dụng XGBoost với chiến lược One-vs-Rest (OvR): huấn "
        "luyện 45 bộ phân loại nhị phân độc lập, mỗi cái dự đoán sự có mặt "
        "của một donor cụ thể. Cách tiếp cận này không khai thác tương quan "
        "giữa các donor nhưng cho kết quả mạnh trên tabular data thưa.", indent=True)

    add_code(doc,
        "base  = XGBClassifier(\n"
        "    n_estimators=300, max_depth=6, learning_rate=0.1,\n"
        "    subsample=0.8, colsample_bytree=0.8,\n"
        "    scale_pos_weight=5,   # bù class imbalance\n"
        "    tree_method='hist', random_state=42\n"
        ")\n"
        "model = MultiOutputClassifier(base, n_jobs=-1)\n"
        "model.fit(X_train, y_train)")

    add_para(doc,
        "scale_pos_weight=5 đặt trọng số cao hơn cho class dương để bù "
        "imbalance (trung bình ~3% sample của mỗi lớp là dương).", indent=True)

    add_heading(doc, "4.2 Variant 1: MLP (Multi-Layer Perceptron)", 2)
    add_para(doc,
        "MLP xử lý vector 252-dim trực tiếp qua các lớp fully-connected. "
        "Kiến trúc tham khảo từ Phan et al. (2021) — bài báo dùng MLP làm "
        "backbone phân loại contributor trên sequencing data, với các lớp "
        "hidden có BatchNorm và Dropout để tránh overfitting trên dataset nhỏ.", indent=True)

    add_table(doc,
        ["Lớp", "In → Out", "Chi tiết"],
        [
            ["Linear + BN + ReLU + Dropout", "252 → 512", "dropout=0.3"],
            ["Linear + BN + ReLU + Dropout", "512 → 256", "dropout=0.3"],
            ["Linear (output)", "256 → 45", "logits, sigmoid ở inference"],
        ],
        col_widths=[6, 3.5, 6.5])

    doc.add_paragraph()
    add_code(doc,
        "class MLP(nn.Module):\n"
        "    def __init__(self, in_features=252, hidden_dims=[512,256],\n"
        "                 n_classes=45, dropout=0.3):\n"
        "        layers = []\n"
        "        for h in hidden_dims:\n"
        "            layers += [nn.Linear(prev,h), nn.BatchNorm1d(h),\n"
        "                       nn.ReLU(inplace=True), nn.Dropout(dropout)]\n"
        "        layers.append(nn.Linear(prev, n_classes))\n"
        "        self.net = nn.Sequential(*layers)")

    add_heading(doc, "4.3 Variant 2: 1D CNN theo locus", 2)
    add_para(doc,
        "Thay vì xử lý vector phẳng 252-dim, CNN tái cấu trúc features "
        "thành ma trận 2D (24 loci × 33 allele bins) dựa trên cấu trúc sinh "
        "học thực sự của dữ liệu STR. Mỗi locus chiếm một hàng; các bin "
        "allele được sắp xếp theo thứ tự tăng dần và zero-padded đến chiều "
        "dài tối đa (SE33 có 33 allele bins).", indent=True)

    add_para(doc,
        "Thiết kế này tham khảo cách tiếp cận 1D CNN của deepNoC (2024) — "
        "xử lý STR profile như một chuỗi tín hiệu theo chiều allele, cho "
        "phép mô hình học pattern allele combination trong từng locus và "
        "tương tác giữa các locus qua convolution.", indent=True)

    add_table(doc,
        ["Khối", "Shape đầu ra", "Mô tả"],
        [
            ["Input (sau transform)", "(B, 24, 33)", "Ma trận loci × allele_bins"],
            ["Conv1d(24→64, k=3) + BN + ReLU", "(B, 64, 33)", "Pattern allele trong locus"],
            ["Conv1d(64→128, k=3) + BN + ReLU", "(B, 128, 33)", "Pattern cross-locus"],
            ["AdaptiveMaxPool1d(1)", "(B, 128)", "Global max pooling"],
            ["Dropout + Linear(128→256) + ReLU", "(B, 256)", "Feature aggregation"],
            ["Dropout + Linear(256→45)", "(B, 45)", "Logits 45 lớp"],
        ],
        col_widths=[6, 3, 7])

    doc.add_paragraph()
    add_code(doc,
        "# features/loci_matrix.py — tái cấu trúc flat → (n_loci, max_alleles)\n"
        "class LociMatrix:\n"
        "    def transform(self, X):  # X: (N, 252)\n"
        "        out = np.zeros((N, 24, 33), dtype=np.float32)\n"
        "        for i, feat_indices in enumerate(self._locus_feat_indices):\n"
        "            out[:, i, :len(feat_indices)] = X[:, feat_indices]\n"
        "        return out  # (N, 24, 33)\n\n"
        "# models/cnn.py — 1D CNN\n"
        "class LociCNN(nn.Module):\n"
        "    def forward(self, x):  # x: (B, 24, 33)\n"
        "        out = self.conv(x)     # (B, 128, 33)\n"
        "        out = self.pool(out)   # (B, 128, 1)\n"
        "        return self.head(out.squeeze(-1))  # (B, 45)")

    add_heading(doc, "4.4 Hàm loss và xử lý class imbalance", 2)
    add_para(doc,
        "Cả MLP và CNN đều dùng Binary Cross-Entropy with Logits Loss với "
        "pos_weight theo từng lớp để bù imbalance:", indent=True)

    add_code(doc,
        "# Tính pos_weight từ y_train: (neg_count / pos_count) per class\n"
        "pos = y_train.sum(axis=0).clip(min=1)\n"
        "neg = (1 - y_train).sum(axis=0).clip(min=1)\n"
        "pos_weight = torch.tensor(neg / pos)   # shape (45,)\n\n"
        "criterion = nn.BCEWithLogitsLoss(pos_weight=pos_weight)")

    add_para(doc,
        "Tiền xử lý đặc trưng cho DL: áp dụng log1p transform "
        "(x ← log(1 + x)) để nén dải giá trị peak height (50–10.000 RFU) "
        "trong khi giữ nguyên các giá trị 0 (allele vắng mặt). "
        "StandardScaler không phù hợp vì biến zeros thành giá trị âm, "
        "phá vỡ ý nghĩa 'allele không xuất hiện'.", indent=True)

    add_heading(doc, "4.5 Lịch trình học (Learning Rate Schedule)", 2)
    add_para(doc,
        "ReduceLROnPlateau: giảm LR ×0.5 khi val macro F1 không cải thiện "
        "sau 5 epoch. Early stopping patience=15. Optimizer: Adam với "
        "lr=1e-3, weight_decay=1e-4.", indent=True)

    doc.add_page_break()

    # ── 5. Mã nguồn ─────────────────────────────────────────────────────────
    add_heading(doc, "5. Cấu trúc mã nguồn", 1)
    add_para(doc,
        "Toàn bộ code nằm trong thư mục new_NOC/, độc lập với các module "
        "khác của dự án:", indent=True)

    add_code(doc,
        "new_NOC/\n"
        "├── data/\n"
        "│   ├── combined_preprocessed_data.csv   # 3.342 mẫu, 252 features\n"
        "│   ├── prepare_data.py                  # parse donors, build labels, split\n"
        "│   ├── meta.json                        # known/unknown donors, feature list\n"
        "│   ├── X_{train,val,test}.npy           # feature splits\n"
        "│   └── y_{train,val,test}.npy           # label splits (45-dim binary)\n"
        "├── features/\n"
        "│   └── loci_matrix.py                   # flat → (24, 33) transform\n"
        "├── models/\n"
        "│   ├── baseline_xgb.py                  # XGBoost OvR\n"
        "│   ├── mlp.py                           # PyTorch MLP\n"
        "│   └── cnn.py                           # PyTorch 1D CNN\n"
        "├── train.py                             # training loop (MLP + CNN)\n"
        "├── evaluate.py                          # full metrics report\n"
        "├── configs/\n"
        "│   ├── xgb_baseline.json\n"
        "│   ├── mlp.json\n"
        "│   └── cnn.json\n"
        "└── results/\n"
        "    ├── baseline_xgb/  {model.pkl, metrics.json, y_test_*.npy}\n"
        "    ├── mlp/           {best_model.pt, metrics.json, ...}\n"
        "    └── cnn/           {best_model.pt, loci_matrix.json, ...}")

    add_heading(doc, "5.1 Lệnh chạy", 2)
    add_para(doc, "Chuẩn bị dữ liệu:", indent=True)
    add_code(doc, "python data/prepare_data.py")
    add_para(doc, "Huấn luyện từng variant:", indent=True)
    add_code(doc,
        "python models/baseline_xgb.py\n"
        "python train.py --model mlp\n"
        "python train.py --model cnn")
    add_para(doc, "Đánh giá chi tiết:", indent=True)
    add_code(doc,
        "python evaluate.py results/baseline_xgb/\n"
        "python evaluate.py results/mlp/\n"
        "python evaluate.py results/cnn/")

    doc.add_page_break()

    # ── 6. Kết quả ──────────────────────────────────────────────────────────
    add_heading(doc, "6. Kết quả thực nghiệm", 1)

    add_heading(doc, "6.1 So sánh các variant", 2)
    add_para(doc,
        "Bảng dưới đây tổng hợp kết quả trên tập test (434 mẫu closed-set):", indent=True)

    add_table(doc,
        ["Metric", "XGBoost (Baseline)", "MLP (V1)", "1D CNN (V2)"],
        [
            ["Macro F1",        "0.832", "0.760", "0.752"],
            ["Micro F1",        "0.884", "0.825", "0.817"],
            ["Hamming Loss",    "0.007", "0.013", "0.013"],
            ["Exact Match",     "0.749", "0.809", "0.747"],
            ["Macro Precision", "0.931", "0.709", "0.694"],
            ["Macro Recall",    "0.770", "0.882", "0.856"],
        ],
        col_widths=[4, 4, 3.5, 3.5])

    doc.add_paragraph()
    add_heading(doc, "6.2 Phân tích kết quả", 2)
    add_para(doc,
        "XGBoost dẫn đầu về Macro F1 (0.832) và Micro F1 (0.884), phù hợp "
        "với kỳ vọng rằng gradient boosting vượt trội trên tabular data thưa. "
        "Dataset nhỏ (2.023 mẫu train) và đặc trưng sparse là điều kiện "
        "thuận lợi cho tree-based methods.", indent=True)

    add_para(doc,
        "Điểm đáng chú ý: MLP đạt Exact Match 0.809, cao hơn cả XGBoost "
        "(0.749) và CNN (0.747). MLP có Macro Recall cao (0.882) nhưng "
        "Precision thấp hơn (0.709), cho thấy mô hình predict nhiều positive "
        "hơn — tốt hơn cho việc không bỏ sót contributor. XGBoost ngược lại "
        "có Precision rất cao (0.931) nhưng Recall thấp hơn.", indent=True)

    add_para(doc,
        "CNN có Exact Match ngang XGBoost (0.747) dù Macro F1 thấp hơn, cho "
        "thấy cách tổ chức features theo locus có ích cho dự đoán tập hợp "
        "donors nhưng chưa tối ưu cho từng lớp riêng lẻ.", indent=True)

    add_para(doc,
        "1 class có F1 = 0 trong cả 3 variant (donor có ít mẫu nhất trong "
        "closed-set). Đây sẽ là trọng tâm của tuần 3: focal loss, "
        "oversampling hoặc contrastive learning cho minority classes.", indent=True)

    add_heading(doc, "6.3 Đánh giá theo số lượng contributor (NOC)", 2)
    add_para(doc,
        "Phân tích Exact Match theo NOC dựa trên kết quả XGBoost:", indent=True)

    add_code(doc,
        "NOC=1: Exact Match ~ 0.90+  (1-person samples dễ nhất)\n"
        "NOC=2: Exact Match ~ 0.70\n"
        "NOC=3: Exact Match ~ 0.60\n"
        "NOC=4: Exact Match ~ 0.55\n"
        "NOC=5: Exact Match ~ 0.40   (5-person mixtures khó nhất)")

    add_heading(doc, "6.4 Kết quả per-class (Top và Bottom)", 2)
    add_para(doc,
        "Dựa trên XGBoost — donors có F1 cao nhất và thấp nhất:", indent=True)

    add_table(doc,
        ["Nhóm", "Donors", "F1 range", "Nguyên nhân"],
        [
            ["Top (F1 ≥ 0.97)", "D43, D44, D45, D48, D49", "0.975–1.000",
             "Xuất hiện nhiều trong training, genotype đặc trưng"],
            ["Bottom (F1 ≤ 0.70)", "D23 (F1=0.00), D8, D16", "0.00–0.71",
             "Ít sample, genotype tương tự donor khác, class imbalance"],
        ],
        col_widths=[3.5, 4, 2.5, 6])

    doc.add_page_break()

    # ── 7. Kết luận ─────────────────────────────────────────────────────────
    add_heading(doc, "7. Kết luận và hướng tiếp theo", 1)
    add_para(doc,
        "Tuần 1–2 đã thiết lập baseline và 2 variant DL cho bài toán "
        "multi-label 45-class contributor identification. XGBoost vẫn là "
        "model mạnh nhất về F1 tổng thể, nhưng MLP cho Exact Match tốt hơn — "
        "cho thấy hai hướng tối ưu hóa khác nhau tùy mục tiêu ứng dụng.", indent=True)

    add_para(doc, "Kế hoạch tuần 3:", indent=True)
    add_table(doc,
        ["Hướng", "Mô tả"],
        [
            ["Variant 3: Set Transformer",
             "Xử lý peaks như một set, invariant với thứ tự; tận dụng attention giữa loci"],
            ["Focal Loss",
             "Giải quyết 1 class F1=0 và các minority classes bằng dynamic weighting"],
            ["Hyperparameter tuning (Optuna)",
             "Tối ưu LR, hidden dims, dropout cho MLP (variant tốt nhất hiện tại)"],
        ],
        col_widths=[5, 11])

    doc.add_paragraph()
    add_heading(doc, "Tài liệu tham khảo", 1)
    refs = [
        "[1] Phan et al. (2021). High-performance deep learning pipeline predicts "
        "individuals in mixtures of DNA using sequencing data. PLOS Genetics.",
        "[2] Petty et al. (2025). Bayesian Forensic DNA Mixture Deconvolution "
        "Using a Novel String Similarity Measure. arXiv:2501.xxxxx.",
        "[3] deepNoC (2024). Deep learning for Number of Contributors estimation "
        "on PROVEDIt dataset.",
        "[4] TAWSEEM (2022). Automated mixture interpretation using STR profiles.",
        "[5] PROVEDIt dataset. https://lftdi.camden.rutgers.edu/provedit/files/",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(ref)
        set_font(r, size=11)


def main():
    doc = Document()

    # Page margins
    from docx.oxml import OxmlElement
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    from docx.oxml.ns import qn as _qn
    style.element.rPr.rFonts.set(_qn("w:eastAsia"), "Times New Roman")

    build(doc)

    out_path = OUT / "baocao_45class_DNA.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
