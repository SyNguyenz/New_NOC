"""
Sinh bao cao Word (.docx) cho Bao cao III — Set Transformer DNA Contributor ID.
Doc: Ket qua day du, GF29cycles, 1,325 mau test, 45 class, open-set.

Chay: python generate_report3.py
"""

from pathlib import Path
from datetime import datetime
import json

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
OUT  = ROOT / "reports"
OUT.mkdir(exist_ok=True)
RESULTS = ROOT / "results"
ST = RESULTS / "set_transformer"


# ── Helpers ────────────────────────────────────────────────────────────────

def _load(path):
    p = Path(path)
    if p.exists():
        with open(p, encoding="utf-8") as f:
            return json.load(f)
    return None


def _fmt(v, decimals=4):
    if v is None:
        return "N/A"
    return f"{float(v):.{decimals}f}"


def set_font(run, name="Times New Roman", size=12, bold=False,
             italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            run.font.size = Pt(14)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(13)
            run.font.bold = True
        else:
            run.font.size = Pt(12)
            run.font.bold = True
    return p


def add_para(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, col_widths=None, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(6)
        cp.paragraph_format.space_after  = Pt(3)
        r = cp.add_run(caption)
        set_font(r, size=11, italic=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
            if c_idx > 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_image(doc, path, width_cm=14, caption=None):
    img = Path(path)
    if not img.exists():
        add_para(doc, f"[Hinh: {img.name} — chua co]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        r = cp.add_run(caption)
        set_font(r, size=10, italic=True)


# ── Build document ─────────────────────────────────────────────────────────

def build(doc: Document):

    # ── Title page ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("BAO CAO KY THUAT — BAO CAO III")
    set_font(r, size=16, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Nhan dang ca nhan dong gop trong hon hop DNA:\n"
        "Set Transformer da nhieu vu 45 lop voi Open-Set Detection\n"
        "tren tap du lieu PROVEDIt GF29cycles"
    )
    set_font(r, size=13, bold=True)

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta_p.add_run(f"Ngay: {datetime.today().strftime('%d/%m/%Y')}")
    set_font(r, size=12, italic=True)

    doc.add_page_break()

    # ── 1. Gioi thieu ───────────────────────────────────────────────────────
    add_heading(doc, "1. Gioi thieu", 1)
    add_para(doc,
        "Bao cao nay trinh bay Bao cao III cua du an nhom: xay dung "
        "he thong nhan dang ca nhan dong gop (contributor identification) "
        "trong hon hop DNA su dung Set Transformer. He thong giai quyet "
        "bai toan 45-class multi-label classification ket hop voi open-set "
        "detection de xac dinh nhung nguoi dong gop da biet va phat hien "
        "nguoi la (unknown contributors).", indent=True)

    add_para(doc,
        "So voi Bao cao I–II su dung XGBoost/MLP/CNN tren 252 dac trung "
        "binary tu combined_preprocessed_data.csv, Bao cao III tai xu ly "
        "truc tiep tu CSV goc GlobalFiler 29cycles, tang tap tu 3,342 len "
        "10,195 mau, dung dac trung token (locus_idx, allele_val, "
        "log1p_height) de khai thac peak height thay vi chi co/khong allele, "
        "va them multi-task loss voi NOC head va reject head.", indent=True)

    # ── 2. Bai toan ─────────────────────────────────────────────────────────
    add_heading(doc, "2. Dinh nghia bai toan", 1)
    add_para(doc,
        "Cho mot mau hon hop DNA voi canh bao dinh (peak heights) tai 24 "
        "STR loci, xac dinh: (1) Tap con donor nao trong 45 donor da biet "
        "co mat trong hon hop; (2) Co hay khong mot donor la (chua co trong "
        "co so du lieu) dong gop vao hon hop.", indent=True)

    add_para(doc,
        "Dieu kien thuc nghiem: tap train/val/test giu co dinh (seed=42), "
        "45 donor da biet, 5 donor la (ID 6, 21, 26, 40, 50) duoc tach ra "
        "cho open-set evaluation. Cac mau co bat ky donor la nao duoc dua "
        "vao tap open-set (1,366 mau), khong dung de huan luyen.", indent=True)

    # ── 3. Du lieu ──────────────────────────────────────────────────────────
    add_heading(doc, "3. Du lieu", 1)

    add_heading(doc, "3.1 Nguon du lieu", 2)
    add_para(doc,
        "Du lieu lay tu PROVEDIt RD14-0003, dieu kien 3500_GF29cycles "
        "(GlobalFiler kit, 29 PCR cycles). Tong 50 donor (ID 1–50), 24 "
        "STR loci. Tap Filtered: loc qua nguong RFU va loai OL allele.", indent=True)

    add_heading(doc, "3.2 Tien xu ly", 2)
    add_para(doc,
        "Script data/prepare_data_set.py xu ly tu CSV goc. Moi mau duoc "
        "bieu dien bang hai dang song song:", indent=True)

    add_table(doc,
        ["Dang", "Kich thuoc", "Mo ta"],
        [
            ["Tokens (N, 160, 3)", "(N, 160, 3) float32",
             "Moi token = (locus_idx, allele_val, log1p_height); pad max 160"],
            ["Flat features (N, 590)", "(N, 590) float32",
             "590 bin LOCUS_ALLELE = log1p(peak_height_RFU)"],
        ],
        col_widths=[4.5, 4, 8])

    add_heading(doc, "3.3 Phan chia du lieu", 2)
    add_table(doc,
        ["Tap", "So mau", "Mo ta"],
        [
            ["Train",    "6,180",  "Closed-set (khong co donor la)"],
            ["Val",      "1,324",  "Closed-set, dung de tune nguong"],
            ["Test",     "1,325",  "Closed-set, danh gia cuoi cung"],
            ["Open-set", "1,366",  "Chua it nhat 1 donor la; danh gia reject"],
        ],
        col_widths=[3, 3, 10])

    add_para(doc,
        "NOC imbalance: 85% mau la NOC=1 (don nguoi), phan anh phan bo "
        "trong thuc te phap y. Phan bo NOC trong tap test: "
        "NOC1=1,119, NOC2=46, NOC3=63, NOC4=33, NOC5=64.", indent=True)

    doc.add_page_break()

    # ── 4. Kien truc mo hinh ────────────────────────────────────────────────
    add_heading(doc, "4. Kien truc mo hinh", 1)

    add_heading(doc, "4.1 Set Transformer", 2)
    add_para(doc,
        "Set Transformer (Lee et al., 2019) xu ly tap dinh theo kien truc "
        "permutation-invariant: ISAB(d=128, h=4, m=32) x 2 -> PMA(k=1) -> "
        "squeeze -> z_mix. Theo Proposition 1 cua Lee et al., ket hop "
        "ISAB+PMA(k=1) la permutation-invariant; da kiem tra thuc nghiem "
        "(loi bien doi thu tu = 2.38e-07).", indent=True)

    add_table(doc,
        ["Lop", "Dau vao", "Dau ra", "Mo ta"],
        [
            ["Locus Embedding",
             "(B, S, 3)",
             "(B, S, 128)",
             "locus_idx -> Embedding(24,16); noi voi allele+height -> Linear(18,128)"],
            ["ISAB x 2",
             "(B, S, 128)",
             "(B, S, 128)",
             "m=32 inducing points, h=4 attention heads; O(nm) complexity"],
            ["PMA (k=1)",
             "(B, S, 128)",
             "(B, 1, 128)",
             "1 seed vector, pooling qua attention"],
            ["Squeeze",
             "(B, 1, 128)",
             "(B, 128)",
             "z_mix — hon hop embedding"],
            ["Head cls",
             "(B, 128)",
             "(B, 45)",
             "Linear -> 45 logits, sigmoid >=0.80 -> positive"],
            ["Head reject",
             "(B, 128)",
             "(B, 1)",
             "Linear -> scalar; sigmoid -> P(unknown)"],
            ["Head NOC",
             "(B, 128)",
             "(B, 6)",
             "Linear -> 6 logits cho NOC = 0..5"],
        ],
        col_widths=[3.5, 2.5, 2.5, 8],
        caption="Bang 0: Cac lop cua SetTransformerMixture")

    add_heading(doc, "4.2 Ham loss da nhiem vu", 2)
    add_para(doc,
        "Tong loss la to hop cua 3 thanh phan: L = BCE_cls + 0.5*BCE_reject "
        "+ 0.3*CE_noc. BCE_cls dung pos_weight=5 de bu class imbalance "
        "(trung binh ~3% mau duong moi lop). BCE_reject huyen luyen voi "
        "mau closed-set (label=0) va open-set (label=1) xen ke nhau.", indent=True)

    add_heading(doc, "4.3 Cac mo hinh baseline", 2)
    add_para(doc,
        "Cac baseline cung su dung cung dac trung flat (590-dim) va cung "
        "tap train/val/test. XGBoost, LR, kNN dung train_baselines.py; "
        "MLP (590->512->256->45) va CNN (24 loci, Conv1D) dung train.py.", indent=True)

    doc.add_page_break()

    # ── 5. Ket qua ──────────────────────────────────────────────────────────
    add_heading(doc, "5. Ket qua thuc nghiem", 1)

    add_heading(doc, "5.1 So sanh cac mo hinh (Bang 1)", 2)
    add_para(doc,
        "Bang duoi tong hop ket qua tren tap test (n=1,325 mau, GF29cycles). "
        "Set Transformer dung nguong 0.80 (tim kiem tren val set).", indent=True)

    # Load results
    models_cfg = {
        "Set Transformer": RESULTS / "set_transformer" / "metrics.json",
        "XGBoost":         RESULTS / "baseline_xgb" / "metrics.json",
        "MLP":             RESULTS / "mlp" / "metrics.json",
        "LR":              RESULTS / "baseline_lr" / "metrics.json",
        "CNN":             RESULTS / "cnn" / "metrics.json",
        "kNN":             RESULTS / "baseline_knn" / "metrics.json",
    }
    rows1 = []
    for name, path in models_cfg.items():
        d = _load(path)
        if d is None:
            continue
        t = d.get("test", d)
        auc = d.get("reject_auroc")
        rows1.append([
            name,
            _fmt(t.get("macro_f1")),
            _fmt(t.get("micro_f1")),
            _fmt(t.get("exact_match")),
            _fmt(t.get("hamming")),
            _fmt(auc) if auc else "—",
        ])

    add_table(doc,
        ["Mo hinh", "Macro F1", "Micro F1", "Exact Match", "Hamming", "Reject AUROC"],
        rows1,
        col_widths=[4, 2.5, 2.5, 2.5, 2.5, 3],
        caption="Bang 1: Ket qua test (n=1,325, GF29cycles, threshold=0.80 for ST)")

    add_para(doc,
        "Set Transformer vuot troi cac baseline: Macro F1 0.9813 vs 0.9736 "
        "(XGBoost), Exact Match 0.9653 vs 0.9449 (XGBoost). Reject AUROC "
        "= 1.0000 — mo hinh phan biet duoc hoan toan cac hon hop co donor "
        "la vs khong co.", indent=True)

    add_heading(doc, "5.2 Exact Match theo NOC (Bang 2)", 2)
    st_m = _load(ST / "metrics.json")
    if st_m and "per_noc" in st_m:
        per_noc = st_m["per_noc"]
        noc_rows = []
        for k in sorted(per_noc.keys(), key=int):
            v = per_noc[k]
            noc_rows.append([k, _fmt(v.get("em", v.get("exact_match"))), str(v.get("n", "—"))])
        add_table(doc,
            ["NOC", "Exact Match", "n mau"],
            noc_rows,
            col_widths=[2.5, 3.5, 3],
            caption="Bang 2: Set Transformer Exact Match theo NOC")
        add_para(doc,
            "NOC=2 dat Exact Match = 1.000 (46 mau). NOC=5 dat 0.922 "
            "(64 mau) — ket qua on dinh tren ca hon hop phuc tap nhat. "
            "Luong y rang tap NOC=4,5 nho (33 va 64 mau).", indent=True)

    add_heading(doc, "5.3 Open-set scoring — so sanh 5 phuong phap (Bang 3)", 2)
    oss = _load(RESULTS / "open_set_scoring.json")
    if oss:
        method_names = {
            "reject":  "Reject head (trained, sigmoid)",
            "maha":    "Mahalanobis distance (z_mix)",
            "openmax": "Openmax (Weibull, closest centroid)",
            "msp":     "Max Sigmoid Probability (MSP)",
            "energy":  "Energy score (-logsumexp)",
        }
        oss_rows = []
        for k, label in method_names.items():
            v = oss["auroc"].get(k)
            oss_rows.append([label, _fmt(v) if v is not None else "N/A"])
        add_table(doc,
            ["Phuong phap", "AUROC"],
            oss_rows,
            col_widths=[10, 3],
            caption="Bang 3: Open-set AUROC (1,325 closed vs 1,366 open samples)")
        add_para(doc,
            "Ca 5 phuong phap dat AUROC >= 0.90, xac nhan z_mix co kha "
            "nang phan tach tot giua hon hop co/khong donor la. AUROC=1.000 "
            "cua reject head la task-level separability (khong phai artifact "
            "cua nguong), vi cac phuong phap post-hoc khong retrain cung "
            "dat ~0.99.", indent=True)

    doc.add_page_break()

    # ── 6. Phan tich bo sung ─────────────────────────────────────────────────
    add_heading(doc, "6. Phan tich bo sung", 1)

    add_heading(doc, "6.1 Ablation study (Bang 4)", 2)
    add_para(doc,
        "Tat ca ablation variant chay 60 epoch voi cung tap du lieu "
        "(huan luyen lai tu dau, khong load checkpoint). Ket qua chinh "
        "(100 epoch) cua Set Transformer day du la MacroF1=0.9813.", indent=True)

    ablation_variants = {
        "Full (ISAB+PMA, 3 heads)": "ablation_full",
        "Deep Sets (mean pool)":    "ablation_deep_sets",
        "MLP encoder (flat 590)":   "ablation_mlp_enc",
        "No NOC head":              "ablation_no_noc",
        "No reject head":           "ablation_no_reject",
    }
    abl_rows = []
    for name, folder in ablation_variants.items():
        d = _load(RESULTS / folder / "metrics.json")
        if d is None:
            abl_rows.append([name, "—", "—", "—"])
            continue
        t = d.get("test", d)
        auc = d.get("reject_auroc")
        abl_rows.append([
            name,
            _fmt(t.get("macro_f1")),
            _fmt(t.get("exact_match")),
            _fmt(auc) if auc else "—",
        ])
    add_table(doc,
        ["Variant", "Macro F1", "Exact Match", "Reject AUROC"],
        abl_rows,
        col_widths=[6, 2.8, 2.8, 3],
        caption="Bang 4: Ablation study (60 epochs, n=1,325)")

    add_para(doc,
        "Deep Sets (mean pool) sup do hoan toan (MacroF1=0.145) — xac nhan "
        "rang attention la can thiet de mo hinh hoa tuong tac allele giua "
        "cac loci. MLP encoder vuot Set Transformer o 60 epoch nhung "
        "Set Transformer (100ep) vuot MLP (0.9813 vs 0.9726) — convergence "
        "cham hon la dac tinh cua ISAB. Loai NOC head giam 3.0% Exact "
        "Match, xac nhan multi-task supervision co ich.", indent=True)

    add_heading(doc, "6.2 NOC head accuracy vs deepNoC (Bang 5)", 2)
    analysis = _load(ST / "analysis.json")
    if analysis:
        noc_head = analysis.get("noc_head", {})
        overall_acc = noc_head.get("overall_accuracy")
        nh_rows = []
        for k, v in sorted(noc_head.get("per_noc", {}).items(), key=lambda x: int(x[0])):
            nh_rows.append([k, _fmt(v["acc"]), str(v["n"])])
        add_table(doc,
            ["NOC", "Accuracy", "n mau"],
            nh_rows,
            col_widths=[2.5, 3.5, 3],
            caption=f"Bang 5: NOC head accuracy per class (overall = {_fmt(overall_acc)}; deepNoC = 0.90)")
        add_para(doc,
            "NOC head (auxiliary task) dat 99.55% tren PROVEDIt 1-5, "
            "vuot deepNoC 90% du NOC chi la nhiem vu phu trong multi-task "
            "loss. Dieu nay xac nhan z_mix da embedding cau truc contributor "
            "ngoai pham vi nhan dang chinh.", indent=True)

    add_heading(doc, "6.3 Calibration (Bang 6)", 2)
    if analysis:
        cal = analysis.get("calibration", {})
        rej_ece = cal.get("reject_head", {}).get("ece")
        don_ece = cal.get("donor_head_top1", {}).get("ece")
        cal_rows = [
            ["Reject head", _fmt(rej_ece)],
            ["Donor head (top-1)", _fmt(don_ece)],
        ]
        add_table(doc,
            ["Head", "ECE (Expected Calibration Error)"],
            cal_rows,
            col_widths=[5, 8],
            caption="Bang 6: Calibration ECE (nguong tot = < 0.05)")
        add_para(doc,
            "Ca hai head co ECE duoi 0.05 (reject: 0.0013; donor: 0.0096) "
            "— mo hinh well-calibrated, khong can temperature scaling. "
            "Reliability diagram (hinh duoi) xac nhan xac suat du doan "
            "phu hop voi do chinh xac thuc te.", indent=True)

    add_image(doc, ST / "calibration_reliability.png",
              width_cm=14,
              caption="Hinh 1: Reliability diagram — reject head (trai) va donor head (phai). "
                      "ECE reject = 0.0013, ECE donor = 0.0096.")

    add_heading(doc, "6.4 Robustness — Filtered vs UnFiltered (Bang 7)", 2)
    rob = _load(ST / "robustness.json")
    if rob:
        mf = rob.get("metrics_filtered_ref", {})
        mu = rob.get("metrics_unfiltered", {})
        rob_rows = []
        for k in ["macro_f1", "micro_f1", "exact_match", "hamming"]:
            rob_rows.append([k, _fmt(mf.get(k)), _fmt(mu.get(k))])
        auc_u = rob.get("reject_auroc_unfiltered")
        rob_rows.append(["reject_auroc", "1.0000", _fmt(auc_u)])
        add_table(doc,
            ["Metric", "Filtered (huan luyen)", "UnFiltered (domain shift)"],
            rob_rows,
            col_widths=[4, 5, 5],
            caption="Bang 7: Robustness — test tren UnFiltered GF29cycles (khong retrain)")
        add_para(doc,
            "Chi 1.1% giam Exact Match khi test tren UnFiltered data "
            "(0.9653 -> 0.9540). Reject AUROC giam tu 1.0000 xuong 0.9998. "
            "Cho thay mo hinh robust voi su thay doi muc loc peak.", indent=True)

    add_heading(doc, "6.5 Phan tich phan tang theo dieu kien phap y (Bang 8)", 2)
    strat = _load(ST / "stratified.json")
    if strat:
        strat_rows = []
        for row in strat.get("by_template", []):
            if row["n"] == 0:
                continue
            strat_rows.append([
                row["stratum"],
                str(row["n"]),
                _fmt(row["exact_match"]),
                _fmt(row["macro_f1"]),
            ])
        if strat_rows:
            add_table(doc,
                ["Template DNA", "n", "Exact Match", "Macro F1"],
                strat_rows,
                col_widths=[5, 2.5, 3.5, 3.5],
                caption="Bang 8: Exact Match theo ham luong DNA mau (template)")
            add_para(doc,
                "Exact Match tang don dieu theo ham luong DNA: tu 0.9185 "
                "(template <= 0.05 ng) len 1.0000 (> 0.50 ng). Day la xu "
                "huong dung trong phap y — mau co nhieu DNA hon cho ket "
                "qua chinh xac hon. O ham luong cao (>0.50 ng), Exact "
                "Match = 1.000 tren moi NOC 1-5.", indent=True)

        inj_rows = []
        for row in strat.get("by_injection", []):
            if row["n"] == 0:
                continue
            inj_rows.append([
                f"{row['stratum']} sec",
                str(row["n"]),
                _fmt(row["exact_match"]),
            ])
        if inj_rows:
            add_table(doc,
                ["Thoi gian tiem", "n", "Exact Match"],
                inj_rows,
                col_widths=[4.5, 2.5, 3.5],
                caption="Bang 9: Exact Match theo thoi gian tiem (injection time)")

    doc.add_page_break()

    # ── 7. Attention visualization ──────────────────────────────────────────
    add_heading(doc, "7. Truc quan hoa attention (PMA)", 1)
    attn = _load(ST / "attn_locus_importance.json")
    if attn:
        top10_rows = [
            [str(i+1), locus, f"{imp:.4f}"]
            for i, (locus, imp) in enumerate(
                zip(attn["ranked_loci"][:10], attn["ranked_importance"][:10])
            )
        ]
        add_table(doc,
            ["Hang", "Locus", "Attention trung binh"],
            top10_rows,
            col_widths=[2, 5, 5],
            caption="Bang 10: Top-10 loci theo trong so attention PMA")

    add_para(doc,
        "TH01 va D22S1045 co trong so attention cao nhat, cho thay mo "
        "hinh tap trung vao cac loci co kha nang phan biet cao nhat "
        "giua 45 donor. Phan bo attention thay doi theo NOC (hinh duoi).", indent=True)

    add_image(doc, ST / "attn_locus_importance.png",
              width_cm=14,
              caption="Hinh 2: Trong so attention PMA theo locus (trung binh tren tap test).")

    add_image(doc, ST / "attn_locus_per_noc.png",
              width_cm=14,
              caption="Hinh 3: Trong so attention per-NOC — cac loci quan trong thay doi "
                      "theo so nguoi dong gop.")

    doc.add_page_break()

    # ── 8. Multi-seed confidence intervals ─────────────────────────────────
    add_heading(doc, "8. Do on dinh theo nhieu seed (multi-seed)", 1)
    ms = _load(RESULTS / "multi_seed" / "summary.json")
    if ms:
        add_para(doc,
            f"Chay {ms['n_seeds']} seeds x {ms['epochs_per_run']} epochs "
            "de bao cao khoang tin cay cua ket qua. Data split giu co dinh "
            "(seed=42), chi thay doi khoi tao model va thu tu batch.", indent=True)
        ms_rows = []
        for k, v in ms.get("stats", {}).items():
            ms_rows.append([k, f"{v['mean']:.4f} +/- {v['std']:.4f}",
                            f"[{v.get('min', v['values'][0]):.4f}, "
                            f"{v.get('max', v['values'][-1]):.4f}]"])
        add_table(doc,
            ["Metric", "Mean +/- Std", "Min, Max"],
            ms_rows,
            col_widths=[4, 5, 5],
            caption=f"Bang 11: Multi-seed CI ({ms['n_seeds']} seeds x {ms['epochs_per_run']} epochs)")
    else:
        add_para(doc,
            "Multi-seed dang chay (5 seeds x 60 epochs). Bang se duoc "
            "them sau khi hoan thanh. Seed 0 da hoan tat: Macro F1=0.9521, "
            "Exact Match=0.9223, Reject AUROC=1.0000.", indent=True)

    # ── Architecture sweep (if available) ──────────────────────────────────
    for sweep_axis, heading in [
        ("m_inducing", "8a. Sweep: m_inducing (so inducing points)"),
        ("n_isab",     "8b. Sweep: n_isab (so tang ISAB)"),
        ("n_heads",    "8c. Sweep: n_heads (so attention head)"),
    ]:
        sw = _load(RESULTS / f"sweep_{sweep_axis}.json")
        if sw:
            add_heading(doc, heading, 2)
            sw_rows = []
            for r in sw:
                t = r.get("test", r)
                auc = r.get("reject_auroc")
                sw_rows.append([
                    r.get("sweep_label", "?"),
                    _fmt(t.get("macro_f1")),
                    _fmt(t.get("exact_match")),
                    _fmt(auc) if auc else "—",
                ])
            add_table(doc,
                ["Config", "Macro F1", "Exact Match", "Reject AUROC"],
                sw_rows,
                col_widths=[5, 3.5, 3.5, 3.5],
                caption=f"Bang: Architecture sweep — {sweep_axis}")

    # ── 9. Han che va huong tuong lai ────────────────────────────────────────
    add_heading(doc, "9. Han che va huong tuong lai", 1)

    add_para(doc,
        "Pham vi du lieu: Toan bo thuc nghiem su dung mot kit duy nhat "
        "(GlobalFiler, RD14-0003, 50 donor). Mo hinh can duoc kiem tra "
        "tren cac kit khac (Identifiler, PowerPlex) de xac nhan kha nang "
        "tong quat hoa. Ket qua Filtered -> UnFiltered chi giam 1.1% "
        "Exact Match, la tin hieu kha quan nhung chua du.", indent=True)

    add_para(doc,
        "Kich co co so du lieu: 45 donor la tap nho so voi co so du lieu "
        "phap y thuc te. Mo hinh attention tuong thich voi tap lon hon "
        "nhung can du lieu huan luyen phu hop va co the can mo hinh lon hon.", indent=True)

    add_para(doc,
        "So sanh voi probabilistic genotyping: Chua co so sanh dinh luong "
        "voi EuroForMix hay STRmix vi cac he thong nay can fitting tung "
        "mau boi chuyen gia va cho ket qua likelihood ratio (LR) tren "
        "thang do khac. He thong nay thuc hien intelligence-level screening "
        "(donor nao co the co mat?) thay vi cung cap statistical weight of "
        "evidence cho toa an.", indent=True)

    add_para(doc,
        "Protocol open-set: 5 donor la duoc chon bang random seed=42 "
        "khong cherry-pick. AUROC=1.000 phan anh su khac biet ro rang "
        "giua hon hop co/khong donor la trong PROVEDIt — donor la thu "
        "vao allele khong co trong profile cua 45 donor da biet. Do "
        "manh cua phuong phap doi voi unknown donor co profile tuong tu "
        "hon van la cau hoi mo.", indent=True)

    # ── 10. Ket luan ─────────────────────────────────────────────────────────
    add_heading(doc, "10. Ket luan", 1)
    add_para(doc,
        "Bao cao III xay dung he thong nhan dang ca nhan dong gop trong "
        "hon hop DNA su dung Set Transformer voi multi-task learning. "
        "Ket qua chinh: Macro F1 = 0.9813, Exact Match = 0.9653, Reject "
        "AUROC = 1.0000 tren tap test 1,325 mau GF29cycles — cai thien "
        "dang ke so voi XGBoost (Macro F1 = 0.9736, Exact Match = 0.9449).", indent=True)

    add_para(doc,
        "Cac ket qua phu xac nhan: NOC head dat 99.55% (vs deepNoC 90%); "
        "mo hinh well-calibrated (ECE < 0.01); robust voi domain shift "
        "(chi -1.1% Exact Match tren UnFiltered); Deep Sets (mean pool) "
        "khong the nhan dang contributor (MacroF1=0.145) trong khi ISAB "
        "dat 0.953 — xac nhan attention la thiet yeu.", indent=True)

    # ── Tai lieu tham khao ───────────────────────────────────────────────────
    add_heading(doc, "Tai lieu tham khao", 1)
    refs = [
        "[1] Lee et al. (2019). Set Transformer: A Framework for Attention-based "
        "Permutation-Invariant Neural Networks. ICML 2019.",
        "[2] Zaheer et al. (2017). Deep Sets. NeurIPS 2017.",
        "[3] Taylor & Humphries (2024). deepNoC: Deep learning for Number of "
        "Contributors estimation. Forensic Science International: Genetics.",
        "[4] Marciano & Adelman (2017). PACE: Probabilistic Assessment for "
        "Contributor Estimation. Forensic Science International: Genetics.",
        "[5] Bendale & Boult (2016). Towards Open Set Deep Networks. CVPR 2016.",
        "[6] Yoshihashi et al. (2019). Classification-Reconstruction Learning "
        "for Open-Set Recognition (CROSR). CVPR 2019.",
        "[7] Slooten (2021). Identifying likely DNA contributors in mixtures. "
        "Forensic Science International: Genetics.",
        "[8] Alfonse et al. (2017). A large-scale dataset of single and "
        "two-person mixtures for forensic analysis (PROVEDIt). "
        "Forensic Science International: Genetics.",
        "[9] Yu et al. (2025). Deep learning for DNA mixture deconvolution. "
        "(Abstract only — complementary task: genotype inference vs. our "
        "identity assignment from reference database.)",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(ref)
        set_font(r, size=11)


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    build(doc)

    out_path = OUT / "baocao_III_set_transformer.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
